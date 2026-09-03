from pathlib import Path
from docx import Document
from pypdf import PdfReader
from app.services.llm_service import extract_from_image_with_vision_llm


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from a DOCX medical report including paragraphs and tables.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    document = Document(str(file_path))
    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    row_text.append(text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n".join(paragraphs)


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from a PDF medical report.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    reader = PdfReader(str(file_path))
    pages_text = []

    for idx, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages_text.append(text.strip())

    combined_text = "\n\n".join(pages_text)

    # If the PDF has no text layer (e.g. it is a scanned document / photocopied report)
    if len(combined_text.strip()) < 30:
        # Check if there are embedded images in the first page to run vision extraction
        try:
            for page in reader.pages[:3]:
                for image_file_object in page.images:
                    temp_img_path = file_path.parent / f"{file_path.stem}_page_img.png"
                    with open(temp_img_path, "wb") as fp:
                        fp.write(image_file_object.data)
                    vision_result = extract_from_image_with_vision_llm(temp_img_path)
                    if temp_img_path.exists():
                        temp_img_path.unlink(missing_ok=True)
                    if vision_result.get("raw_text"):
                        return vision_result["raw_text"]
        except Exception as err:
            print(f"[PDF Image Extraction Fallback]: {err}")

    return combined_text


def extract_text_from_image(file_path: Path) -> str:
    """
    Extract text from a medical image (PNG, JPG, JPEG, WEBP) using multimodal vision.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    vision_data = extract_from_image_with_vision_llm(file_path)
    return vision_data.get("raw_text", "")


def extract_text_from_file(file_path: Path) -> str:
    """
    Unified text extraction dispatcher for DOCX, PDF, and image files.
    """
    ext = file_path.suffix.lower()

    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in {".png", ".jpg", ".jpeg", ".webp"}:
        return extract_text_from_image(file_path)
    elif ext == ".txt":
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        return file_path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file format for extraction: {ext}")
